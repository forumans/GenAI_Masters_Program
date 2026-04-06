"""
Module 1: Introduction to Text Mining and NLP
Case Study 01
"""

import os
import shutil
import pandas as pd
from docx import Document
from docx.oxml.ns import qn


# =============================================================================
# Question 1: Directory operations
# =============================================================================

# Print current working directory
print("Current working directory:", os.getcwd())

# Create a new directory on the Desktop named "Text Mining and NLP"
desktop = os.path.join(os.path.expanduser("~"), "Desktop")
new_dir = os.path.join(desktop, "Text Mining and NLP")
os.makedirs(new_dir, exist_ok=True)
print("Created directory:", new_dir)

# Make "Text Mining and NLP" the current working directory
os.chdir(new_dir)
print("New current working directory:", os.getcwd())


# =============================================================================
# Question 2: File operations
# =============================================================================

# Create "Greetings.txt" in the current directory and write content
greetings_path = os.path.join(os.getcwd(), "Greetings.txt")
with open(greetings_path, "w") as f:
    f.write("Welcome to Text Mining and Natural Language Processing")
print("Created file:", greetings_path)

# The file is already inside "Text Mining and NLP", so just rename it
welcome_path = os.path.join(os.getcwd(), "Welcome.txt")
if os.path.exists(welcome_path):
    os.remove(welcome_path)
os.rename(greetings_path, welcome_path)
print("Renamed to:", welcome_path)


# =============================================================================
# Question 3: DOCX operations
# =============================================================================

# Path to the provided NLP.docx (in the data folder relative to the project)
script_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(script_dir)
nlp_docx_path = os.path.join(project_root, "data", "01_Text_Mining_And_NLP_Case_Study_01", "NLP.docx")

doc = Document(nlp_docx_path)

# How many paragraphs are in the file?
paragraphs = doc.paragraphs
print("\nNumber of paragraphs in NLP.docx:", len(paragraphs))

# For each paragraph, print the number of words it contains
for i, para in enumerate(paragraphs):
    word_count = len(para.text.split())
    print(f"  Paragraph {i + 1}: {word_count} word(s)")

# Read content from Welcome.txt and insert it as the first paragraph in NLP.docx
with open(welcome_path, "r") as f:
    welcome_content = f.read()

# Insert a new paragraph at the beginning of the document body
first_para = doc.paragraphs[0]._element
new_para = doc.add_paragraph(welcome_content)
# Move the new paragraph's XML element to before the first existing paragraph
first_para.addprevious(new_para._element)

doc.save(nlp_docx_path)
print("\nInserted Welcome.txt content as first paragraph in NLP.docx")
print("Updated NLP.docx saved.")


# =============================================================================
# Question 4: CSV operations with pandas
# =============================================================================

csv_path = os.path.join(project_root, "data", "01_Text_Mining_And_NLP_Case_Study_01", "Employee_Data.csv")

# Read the existing Employee_Data.csv
df = pd.read_csv(csv_path)
print("\nOriginal DataFrame:")
print(df)

# Increment "Salary" by 10%
df["Salary"] = df["Salary"] * 1.10

print("\nUpdated DataFrame:")
print(df)

# Save changes back to Employee_Data.csv
df.to_csv(csv_path, index=False)
print("\nSaved updated data to:", csv_path)
